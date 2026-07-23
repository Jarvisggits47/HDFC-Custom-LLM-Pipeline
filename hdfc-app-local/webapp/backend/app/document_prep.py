"""
Document preparation: PII scanning/redaction, page-aware extraction, and
semantic chunking.

Upgraded from the original sliding-window chunker to the paragraph/sentence
-aware "Tier 1" chunker validated in the Colab notebook — chunks now always
end on a real sentence boundary (never cut a fact in half), and PDF
extraction keeps page numbers so citations can point to a real page, not
just a filename.
"""
import hashlib
import re

# Deliberately conservative patterns — false positives (over-redacting) are
# safer than false negatives in a banking context.
PII_PATTERNS = {
    "pan": re.compile(r"\b[A-Z]{5}[0-9]{4}[A-Z]\b"),
    "aadhaar": re.compile(r"\b\d{4}\s?\d{4}\s?\d{4}\b"),
    "phone": re.compile(r"\b(?:\+91[-\s]?)?[6-9]\d{9}\b"),
    "email": re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b"),
    "account_number": re.compile(r"\b\d{9,18}\b"),
    "card_number": re.compile(r"\b(?:\d[ -]*?){13,16}\b"),
}


def redact_pii(text: str) -> tuple[str, bool]:
    """Returns (redacted_text, was_anything_found)."""
    found = False
    redacted = text
    for label, pattern in PII_PATTERNS.items():
        def _sub(m, label=label):
            nonlocal found
            found = True
            return f"[REDACTED_{label.upper()}]"
        redacted = pattern.sub(_sub, redacted)
    return redacted, found


# --------------------------------------------------------------- extraction
def extract_pdf_pages(file_bytes: bytes) -> list[tuple[int, str]]:
    """Returns [(page_number, text), ...] — 1-indexed pages, so citations
    can report a real page number."""
    from pypdf import PdfReader
    import io
    reader = PdfReader(io.BytesIO(file_bytes))
    return [(i + 1, page.extract_text() or "") for i, page in enumerate(reader.pages)]


def extract_pdf_text(file_bytes: bytes) -> str:
    """Kept for backward compatibility with any caller that just wants the
    flat text (e.g. quick sanity checks)."""
    return "\n".join(t for _, t in extract_pdf_pages(file_bytes))


def extract_docx_pages(file_bytes: bytes) -> list[tuple[int, str]]:
    """Extract text from a .docx file. DOCX has no formal page structure in
    its XML, so the whole document is returned as a single page-1 block.
    Falls back to a clear error if python-docx is not installed."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError(
            "python-docx is not installed — run: pip install python-docx"
        )
    import io
    doc = Document(io.BytesIO(file_bytes))
    # Include table cells as well as body paragraphs so tabular policy data
    # isn't silently dropped.
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = "  ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return [(1, "\n".join(parts))]


# ----------------------------------------------------------------- chunking
_CLEAN_RE = re.compile(r"[^\x09\x0A\x0D\x20-\x7E]")
_WS_RE = re.compile(r"[ \t]+")
_SENTENCE_SPLIT = re.compile(r"(?<=[.!?])\s+(?=[A-Z0-9])")

# Matches "4.1 Introduction", "2.3.1 Eligibility" — numbered section headings.
_SECTION_NUM_RE = re.compile(r"^\d+(?:\.\d+)+\.?\s+[A-Z]")


def _clean_text(text: str) -> str:
    text = _CLEAN_RE.sub(" ", text)
    text = _WS_RE.sub(" ", text)
    return text.strip()


def _split_into_sentences(paragraph: str) -> list[str]:
    return [s.strip() for s in _SENTENCE_SPLIT.split(paragraph) if s.strip()]


def semantic_chunk_page(text: str, target_chars: int = 700, max_chars: int = 900,
                         overlap_sentences: int = 1) -> list[str]:
    """Paragraph-aware, sentence-boundary-respecting chunking. Never cuts a
    sentence in half. Returns a list of chunk strings."""
    text = _clean_text(text)
    if not text:
        return []
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n|\n(?=[A-Z][a-z]+ ?\d*\.?\s*$)", text) if p.strip()]
    if not paragraphs:
        paragraphs = [text]

    chunks: list[str] = []
    current_sentences: list[str] = []
    current_len = 0

    def flush():
        if current_sentences:
            chunks.append(" ".join(current_sentences).strip())

    for para in paragraphs:
        sentences = _split_into_sentences(para) or [para]
        for sent in sentences:
            if current_len + len(sent) > max_chars and current_sentences:
                flush()
                carry = current_sentences[-overlap_sentences:] if overlap_sentences else []
                current_sentences = list(carry)
                current_len = sum(len(s) for s in current_sentences)
            current_sentences.append(sent)
            current_len += len(sent)
            if current_len >= target_chars:
                flush()
                current_sentences, current_len = [], 0
    flush()
    return chunks


def is_probably_real_text(chunk: str, min_alpha_ratio: float = 0.55, min_words: int = 8) -> bool:
    letters = sum(c.isalpha() for c in chunk)
    if len(chunk) == 0 or letters / len(chunk) < min_alpha_ratio:
        return False
    if len(chunk.split()) < min_words:
        return False
    return True


def _is_structural_header(chunk: str) -> bool:
    """Returns True for short (≤15-word) chunks that are document structural
    headings — chapter/section labels, numbered-section titles, TOC entries.
    These chunks reach the retriever but add noise without useful policy content
    (e.g. 'Chapter 4 Current Accounts 4.1 Introduction').

    Three signals, in order of specificity:
      1. Starts with a section keyword + number: 'Chapter 4', 'Section 2.1'
      2. Starts with a dotted-section number:    '4.1 Introduction'
      3. Short line, ≥80 % Title-Case/numeric words, no sentence-ending punctuation
    """
    words = chunk.split()
    if not words or len(words) > 15:
        return False
    stripped = chunk.strip()

    # Signal 1 — explicit label: "Chapter 4 ...", "Section 2.1 ...", "Appendix A ..."
    if re.match(
        r"^(?:chapter|section|part|article|appendix|clause|schedule)\s+[\dA-Z]",
        stripped,
        re.IGNORECASE,
    ):
        return True

    # Signal 2 — numbered section heading: "4.1 Introduction", "2.3.1 Eligibility"
    if _SECTION_NUM_RE.match(stripped):
        return True

    # Signal 3 — high Title-Case ratio, no terminal punctuation.
    # ≥80 % threshold avoids filtering legitimate short factual sentences that
    # start with a capital (e.g. "The minimum balance is Rs 10,000 per month").
    title_or_num = sum(1 for w in words if w and (w[0].isupper() or w[0].isdigit()))
    if len(words) <= 12 and title_or_num / len(words) >= 0.80 and stripped[-1:] not in ".!?":
        return True

    return False


def _chunk_hash(text: str) -> str:
    return hashlib.sha256(text.encode()).hexdigest()[:32]


def chunk_pages(pages: list[tuple[int, str]]) -> list[dict]:
    """Turns [(page_num, text), ...] into a flat list of chunk dicts:
    {text, page, chunk_index, content_hash}. Drops garbage/near-empty chunks
    and structural heading lines (chapter/section headers add retrieval noise)."""
    out = []
    idx = 0
    for page_num, page_text in pages:
        redacted_text, _ = redact_pii(page_text)
        for c in semantic_chunk_page(redacted_text):
            if is_probably_real_text(c) and not _is_structural_header(c):
                out.append({"text": c, "page": page_num, "chunk_index": idx,
                             "content_hash": _chunk_hash(c)})
                idx += 1
    return out


def chunk_text(text: str) -> list[dict]:
    """Same semantic chunker for pasted text with no page structure —
    everything reported as page 1."""
    redacted_text, _ = redact_pii(text)
    out = []
    for i, c in enumerate(semantic_chunk_page(redacted_text)):
        if is_probably_real_text(c) and not _is_structural_header(c):
            out.append({"text": c, "page": 1, "chunk_index": i,
                        "content_hash": _chunk_hash(c)})
    return out
